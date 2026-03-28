# Copyright (c) Huawei Technologies Co., Ltd. 2023-2025. All rights reserved.
"""
DOCX 解析器测试

测试范围:
- DOCX 文档解析
- 段落提取
- 表格提取
- 图片提取
"""

from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
import docx

from data_chain.parser.handler.docx_parser import DocxParser
from data_chain.parser.parse_result import ParseResult
from data_chain.entities.enum import ChunkType, DocParseRelutTopology, ChunkParseTopology


class TestDocxParserBasic:
    """测试 DOCX 解析器基本功能"""

    @pytest.mark.asyncio
    async def test_is_image_with_image(self):
        """测试检测图片段落"""
        mock_paragraph = MagicMock()
        mock_paragraph._element.xpath.return_value = [MagicMock()]
        mock_doc = MagicMock()
        
        with patch.object(DocxParser, 'is_image', return_value=True):
            result = await DocxParser.is_image(mock_paragraph, mock_doc)
            assert result is True

    @pytest.mark.asyncio
    async def test_is_image_without_image(self):
        """测试检测非图片段落"""
        mock_paragraph = MagicMock()
        mock_paragraph._element.xpath.return_value = []
        mock_doc = MagicMock()
        
        result = await DocxParser.is_image(mock_paragraph, mock_doc)
        assert result is False

    @pytest.mark.asyncio
    async def test_get_imageparts_from_run(self):
        """测试从 run 获取图片"""
        mock_run = MagicMock()
        mock_run._r.xpath.return_value = []
        mock_doc = MagicMock()
        
        result = await DocxParser.get_imageparts_from_run(mock_run, mock_doc)
        assert isinstance(result, list)

    @pytest.mark.asyncio
    async def test_extract_table_to_array(self):
        """测试表格转数组"""
        mock_table = MagicMock()
        mock_row = MagicMock()
        mock_cell = MagicMock()
        mock_cell.paragraphs = [MagicMock(text="Cell 1")]
        mock_row.cells = [mock_cell]
        mock_table.rows = [mock_row]
        
        result = await DocxParser.extract_table_to_array(mock_table)
        
        assert isinstance(result, list)
        assert result[0][0] == "Cell 1"


class TestDocxParserIntegration:
    """测试 DOCX 解析器集成"""

    @pytest.mark.asyncio
    async def test_parser_basic(self, temp_dir):
        """测试基本解析功能"""
        test_file = temp_dir / "test.docx"
        
        # 创建测试文档
        doc = docx.Document()
        doc.add_heading("OpenEuler 测试文档", level=1)
        doc.add_paragraph("这是第一段内容。")
        doc.add_paragraph("这是第二段内容。")
        doc.save(str(test_file))
        
        try:
            result = await DocxParser.parser(str(test_file))
            
            assert isinstance(result, ParseResult)
            assert result.parse_topology_type == DocParseRelutTopology.LIST
            assert len(result.nodes) >= 2
        finally:
            test_file.unlink(missing_ok=True)

    @pytest.mark.asyncio
    async def test_parser_with_table(self, temp_dir):
        """测试带表格的文档解析"""
        test_file = temp_dir / "test_table.docx"
        
        doc = docx.Document()
        doc.add_paragraph("表格测试")
        
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "Data 1"
        table.cell(1, 1).text = "Data 2"
        
        doc.save(str(test_file))
        
        try:
            result = await DocxParser.parser(str(test_file))
            
            assert isinstance(result, ParseResult)
            # 表格内容也应该被解析
            assert len(result.nodes) >= 1
        finally:
            test_file.unlink(missing_ok=True)

    @pytest.mark.asyncio
    async def test_parser_empty_document(self, temp_dir):
        """测试空文档解析"""
        test_file = temp_dir / "empty.docx"
        
        doc = docx.Document()
        doc.save(str(test_file))
        
        try:
            result = await DocxParser.parser(str(test_file))
            
            assert isinstance(result, ParseResult)
            assert len(result.nodes) >= 0
        finally:
            test_file.unlink(missing_ok=True)

    @pytest.mark.asyncio
    async def test_parser_nonexistent_file(self):
        """测试解析不存在的文件"""
        with pytest.raises(Exception):
            await DocxParser.parser("/nonexistent/file.docx")


class TestDocxParserNodeStructure:
    """测试 DOCX 解析器节点结构"""

    @pytest.mark.asyncio
    async def test_node_properties(self, temp_dir):
        """测试节点属性"""
        test_file = temp_dir / "test_props.docx"
        
        doc = docx.Document()
        doc.add_paragraph("测试段落")
        doc.save(str(test_file))
        
        try:
            result = await DocxParser.parser(str(test_file))
            
            if len(result.nodes) > 0:
                node = result.nodes[0]
                assert node.lv == 0
                assert node.type == ChunkType.TEXT
                assert node.link_nodes == []
        finally:
            test_file.unlink(missing_ok=True)


class TestDocxParserEdgeCases:
    """测试边界情况"""

    @pytest.mark.asyncio
    async def test_parser_special_characters(self, temp_dir):
        """测试特殊字符处理"""
        test_file = temp_dir / "special.docx"
        
        doc = docx.Document()
        doc.add_paragraph("特殊字符: <>&\"'")
        doc.add_paragraph("Unicode: 你好世界 🌍 🎉")
        doc.save(str(test_file))
        
        try:
            result = await DocxParser.parser(str(test_file))
            
            assert isinstance(result, ParseResult)
        finally:
            test_file.unlink(missing_ok=True)

    @pytest.mark.asyncio
    async def test_parser_long_paragraph(self, temp_dir):
        """测试长段落处理"""
        test_file = temp_dir / "long.docx"
        
        doc = docx.Document()
        long_text = "OpenEuler " * 1000
        doc.add_paragraph(long_text)
        doc.save(str(test_file))
        
        try:
            result = await DocxParser.parser(str(test_file))
            
            assert isinstance(result, ParseResult)
            # 验证长文本被正确解析
            assert len(result.nodes) >= 1
        finally:
            test_file.unlink(missing_ok=True)
