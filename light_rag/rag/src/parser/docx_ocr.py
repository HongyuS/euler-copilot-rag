"""
DOCX 文件解析器（带 OCR 功能）
提取文本内容并对文档中的图片进行 OCR 识别
"""
import os
import io
import logging
import tempfile
import asyncio
from typing import Optional
from docx import Document as DocxDocument
from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.parts.image import ImagePart
from docx.table import _Cell, Table
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.shape import CT_Picture
from PIL import Image

from parser.tools.ocr_tool import OcrTool

logger = logging.getLogger(__name__)


class DocxOcrParser:
    """DOCX 文档解析器（带 OCR 功能）"""
    
    @staticmethod
    def _is_image(graph: Paragraph, doc: Document) -> bool:
        """检查段落是否包含图片"""
        images = graph._element.xpath('.//pic:pic')
        for image in images:
            for img_id in image.xpath('.//a:blip/@r:embed'):
                part = doc.part.related_parts[img_id]
                if isinstance(part, ImagePart):
                    return True
        return False
    
    @staticmethod
    def _get_imageparts_from_run(run, doc: Document) -> list:
        """获取 run 中的所有图片"""
        image_parts = []
        drawings = run._r.xpath('.//w:drawing')
        for drawing in drawings:
            for img_id in drawing.xpath('.//a:blip/@r:embed'):
                part = doc.part.related_parts[img_id]
                if isinstance(part, ImagePart):
                    image_parts.append(part)
        return image_parts
    
    @staticmethod
    def _extract_table_to_array(table: Table) -> list:
        """提取表格为数组"""
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = ''.join([p.text for p in cell.paragraphs])
                row_data.append(cell_text)
            table_data.append(row_data)
        return table_data
    
    @staticmethod
    def _extract_images_from_docx(doc: DocxDocument) -> list:
        """
        从 DOCX 文档中提取所有图片
        
        :param doc: DOCX 文档对象
        :return: 图片数据列表，每个元素是 (图片数据, 图片索引)
        """
        images = []
        try:
            # 遍历文档中的所有关系
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_part = rel.target_part
                        image_bytes = image_part.blob
                        images.append(image_bytes)
                    except Exception as e:
                        logger.warning(f"[DocxOcrParser] 提取图片失败: {e}")
        except Exception as e:
            logger.warning(f"[DocxOcrParser] 遍历文档关系失败: {e}")
        
        return images

    @staticmethod
    async def _ocr_image_bytes(image_bytes: bytes, image_index: int) -> str:
        """
        对图片字节数据进行 OCR 识别
        
        :param image_bytes: 图片字节数据
        :param image_index: 图片索引
        :return: OCR 识别的文本
        """
        temp_file = None
        try:
            # 创建临时文件
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                temp_file = tmp.name
                tmp.write(image_bytes)
            
            # 使用 OCR 工具识别
            ocr_text = await OcrTool.image_to_text(temp_file)
            
            if ocr_text and ocr_text.strip():
                return f"[图片{image_index + 1} OCR内容]:\n{ocr_text}"
            return ""
        except Exception as e:
            logger.warning(f"[DocxOcrParser] 图片 {image_index + 1} OCR 识别失败: {e}")
            return ""
        finally:
            # 清理临时文件
            if temp_file and os.path.exists(temp_file):
                try:
                    os.remove(temp_file)
                except Exception as e:
                    logger.warning(f"[DocxOcrParser] 删除临时文件失败: {e}")
    
    @staticmethod
    async def parse(file_path: str) -> Optional[str]:
        """
        解析 DOCX 文件，包括文本内容和图片 OCR
        
        :param file_path: 文件路径
        :return: 文件内容（包含文本和 OCR 结果）
        """
        try:
            doc = DocxDocument(file_path)
            if not doc:
                logger.error("[DocxOcrParser] 无法打开docx文件")
                return None
            
            paragraphs = []
            
            # 提取文本内容
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            # 提取图片并进行 OCR
            try:
                image_bytes_list = DocxOcrParser._extract_images_from_docx(doc)
                if image_bytes_list:
                    logger.info(f"[DocxOcrParser] 找到 {len(image_bytes_list)} 张图片，开始 OCR 识别")
                    
                    # 并发处理所有图片的 OCR
                    ocr_tasks = [
                        DocxOcrParser._ocr_image_bytes(image_bytes, idx)
                        for idx, image_bytes in enumerate(image_bytes_list)
                    ]
                    ocr_results = await asyncio.gather(*ocr_tasks, return_exceptions=True)
                    
                    # 添加 OCR 结果
                    for idx, ocr_result in enumerate(ocr_results):
                        if isinstance(ocr_result, Exception):
                            logger.warning(f"[DocxOcrParser] 图片 {idx + 1} OCR 处理异常: {ocr_result}")
                            continue
                        if ocr_result and ocr_result.strip():
                            paragraphs.append(ocr_result)
            except Exception as e:
                logger.warning(f"[DocxOcrParser] 图片 OCR 处理失败: {e}")
            
            content = '\n'.join(paragraphs)
            return content if content.strip() else None
        except Exception as e:
            logger.exception(f"[DocxOcrParser] 解析DOCX文件失败: {e}")
            return None
