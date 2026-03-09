"""
DOCX 文件解析器（纯文本，无 OCR）
"""
import logging
from typing import Optional
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class DocxSimpleParser:
    """DOCX 文档解析器（仅提取文本）"""

    @staticmethod
    def parse(file_path: str) -> Optional[str]:
        """
        解析 DOCX 文件，提取文本内容
        :param file_path: 文件路径
        :return: 文件内容
        """
        try:
            doc = DocxDocument(file_path)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    parts.append("\t".join(row_text))
            return "\n".join(parts) if parts else None
        except Exception as e:
            logger.exception(f"[DocxSimpleParser] 解析失败: {file_path}, {e}")
            return None
